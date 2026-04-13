import io
import os,hashlib
import shutil
import tempfile
from typing import Optional, List

from docxtpl import DocxTemplate
from fastapi import UploadFile
from langchain_core.documents import Document
from utils.logger_handler import logger
from langchain_community.document_loaders import PyPDFLoader,TextLoader,UnstructuredPDFLoader

def get_file_md5_hex(file: UploadFile) -> str:
    """计算文件的md5值"""
    md5_obj = hashlib.md5()  # 创建md5对象
    chunk_size = 4096
    try:
        while chunk := file.file.read(chunk_size):
            md5_obj.update(chunk)

        md5_hex = md5_obj.hexdigest()
        return md5_hex
    except Exception as e:
        logger.error(f'[md5计算]文件{file.filename}计算失败,{e}')
    finally:
        # 重置文件指针，以便后续读取
        if hasattr(file.file, 'seek'):
            file.file.seek(0)



def listdir_with_allowed_type(path: str, allowed_type: list[str] = ['.pdf', '.txt']) -> object:
    """查找文件夹内符合文件类型的文件"""
    file_list = []
    if not os.path.isdir(path):
        logger.error(f'[listdir_with_allowed_type]路径{path}不是文件夹')
        return file_list

    for file in os.listdir(path):
        if file.endswith(allowed_type):
            file_list.append(os.path.join(path,file))
    return file_list

def pdf_loader(file: UploadFile, password: Optional[str] = None) -> List[Document]:
    temp_file_path = None
    try:
        # 1. 确保源文件指针在开头（双重保险）
        if hasattr(file.file, 'seek'):
            file.file.seek(0)
        # 2. 创建带.pdf后缀的临时文件（delete=False 保证 PyPDFLoader 能访问）
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as temp_file:
            temp_file_path = temp_file.name
            # 3. 重置指针并流式复制（避免大文件内存溢出）
            if hasattr(file.file, 'seek'):
                file.file.seek(0)
            shutil.copyfileobj(file.file, temp_file)  # 流式复制，高效安全

        # 4. 用临时文件路径加载（PyPDFLoader 仅接受路径字符串）
        loader = UnstructuredPDFLoader(temp_file_path, password=password)
        # loader = PyPDFLoader(temp_file_path, password=password)
        return loader.load()

    except Exception as e:
        logger.error(f"[pdf_loader] 加载文件 {file.filename} 失败: {e}")
        return []

    finally:
        # 5. 安全清理临时文件（无论成功/失败）
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except OSError as cleanup_err:
                logger.warning(f"[pdf_loader] 临时文件清理失败 ({temp_file_path}): {cleanup_err}")



def txt_loader(file: UploadFile):
    # return TextLoader(file_path=, encoding="utf-8").load()
    """从 UploadFile 读取文本并返回 Document 列表"""
    try:
        raw = file.file.read()
        if isinstance(raw, bytes):
            text = raw.decode('utf-8', errors='replace')
        else:
            text = raw
        doc = Document(page_content=text, metadata={'source': file.filename or 'upload'})
        return [doc]
    except Exception as e:
        logger.error(f'[txt_loader] 读取文件{file.filename}失败, {e}')
        return []


def doc_loader(file: UploadFile) -> List[Document]:
    """
    加载 Word 文档 (.doc/.docx) 并返回 LangChain Document 列表
    """
    try:
        # 获取文件名并转换为小写，以便进行不区分大小写的比较
        filename = file.filename.lower()
        
        if not (filename.endswith('.docx') or filename.endswith('.doc')):
            logger.error(f"[doc_loader] 不支持的文件类型：{file.filename}")
            return []
        
        # 将文件指针移回开头，确保从头开始读取
        if hasattr(file.file, 'seek'):
            file.file.seek(0)
        
        # 读取文件内容到 BytesIO
        file_content = file.file.read()
        doc_stream = io.BytesIO(file_content)
        
        # 直接使用 python-docx 加载，避免使用 DocxTemplate
        from docx import Document as DocxDocument
        
        try:
            docx_obj = DocxDocument(doc_stream)
        except Exception as docx_err:
            logger.error(f"[doc_loader] python-docx 加载失败：{docx_err}")
            
            # 如果直接加载失败，尝试用 DocxTemplate
            try:
                doc_stream.seek(0)
                doc_template = DocxTemplate(doc_stream)
                docx_obj = doc_template.docx
                if docx_obj is None:
                    logger.error(f"[doc_loader] DocxTemplate 转换后 docx 为 None")
                    return []
            except Exception as template_err:
                logger.error(f"[doc_loader] DocxTemplate 也失败：{template_err}")
                return []
        
        full_text = []
        
        # 提取所有段落的文本
        if hasattr(docx_obj, 'paragraphs'):
            for paragraph in docx_obj.paragraphs:
                text = paragraph.text.strip()
                if text:  # 只保留非空文本
                    full_text.append(text)
        
        # 如果没有提取到任何文本，尝试提取表格内容
        if not full_text and hasattr(docx_obj, 'tables'):
            for table in docx_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            full_text.append(cell_text)
        
        if not full_text:
            logger.warning(f"[doc_loader] 文件 {file.filename} 中未提取到有效文本内容")
            return []
        
        # 创建 LangChain Document 对象
        document = Document(
            page_content="\n".join(full_text),
            metadata={
                'source': file.filename,
                'file_type': 'docx' if filename.endswith('.docx') else 'doc'
            }
        )
        
        logger.info(f"[doc_loader] 成功从文件 {file.filename} 提取 {len(full_text)} 个段落")
        return [document]
    
    except Exception as e:
        logger.error(f"[doc_loader] 加载文件 {file.filename} 失败：{e}")
        import traceback
        traceback.print_exc()
        return []
